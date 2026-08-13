"""Extração local de conteúdo para revisão de documentos comerciais variados."""

from pathlib import Path
import re


def _resumo(texto: str, formato: str) -> dict:
    texto = "\n".join(linha.strip() for linha in texto.splitlines() if linha.strip())
    valores = sorted(set(re.findall(r"R\$\s*[\d.]+(?:,\d{1,4})?", texto, flags=re.I)))[:100]
    ceps = sorted(set(re.findall(r"\b\d{5}-?\d{3}\b", texto)))[:100]
    prazos = sorted(set(re.findall(r"\b\d+\s+dias?\s+(?:úteis|uteis|corridos)\b", texto, flags=re.I)))[:100]
    return {
        "formato": "documento_generico_v1",
        "tipo_documento": formato,
        "texto_extraido": texto[:100_000],
        "valores_detectados": valores,
        "ceps_detectados": ceps,
        "prazos_detectados": prazos,
        "requer_mapeamento_tarifario": True,
    }


def extrair_documento_generico(caminho: Path, tipo: str) -> dict:
    if tipo == "pdf":
        from pypdf import PdfReader
        texto = "\n".join(pagina.extract_text() or "" for pagina in PdfReader(str(caminho)).pages)
    elif tipo == "docx":
        from docx import Document
        documento = Document(str(caminho))
        partes = [p.text for p in documento.paragraphs]
        partes.extend(" | ".join(c.text for c in linha.cells) for tabela in documento.tables for linha in tabela.rows)
        texto = "\n".join(partes)
    elif tipo in {"png", "jpg", "jpeg"}:
        import pytesseract
        from PIL import Image
        texto = pytesseract.image_to_string(Image.open(caminho), lang="por")
    elif tipo in {"xlsx", "xlsm"}:
        from openpyxl import load_workbook
        workbook = load_workbook(caminho, data_only=True, read_only=True)
        partes = []
        for aba in workbook.worksheets:
            partes.append(f"### {aba.title}")
            for linha in aba.iter_rows(values_only=True):
                valores = [str(valor) for valor in linha if valor is not None]
                if valores:
                    partes.append(" | ".join(valores))
                if sum(len(item) for item in partes) > 100_000:
                    break
        workbook.close()
        texto = "\n".join(partes)
    elif tipo == "xls":
        import xlrd
        workbook = xlrd.open_workbook(str(caminho))
        partes = []
        for aba in workbook.sheets():
            partes.append(f"### {aba.name}")
            for indice in range(aba.nrows):
                valores = [str(valor) for valor in aba.row_values(indice) if valor not in (None, "")]
                if valores:
                    partes.append(" | ".join(valores))
        texto = "\n".join(partes)
    else:
        raise ValueError(f"Extração de .{tipo} ainda não disponível")
    if not texto.strip():
        raise ValueError("Nenhum texto legível foi encontrado no documento")
    return _resumo(texto, tipo)
